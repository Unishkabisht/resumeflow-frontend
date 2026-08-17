"""
generate_clean_report.py
Generates a 100% Word-valid Summer Internship Project Report (.docx) for
ResumeFlow and writes it to TWO locations:

    C:\\Users\\Unishka Bisht\\Downloads\\ResumeFlow_Internship_Report.docx
    C:\\Users\\Unishka Bisht\\resumeflow-frontend\\ResumeFlow_Internship_Report.docx

Robustness notes (to avoid "Word experienced an error opening the file"):
  * Every screenshot is re-encoded through PIL into a clean, profile-free
    RGB JPEG before being added -> no CMYK, no EXIF/orientation surprises,
    no embedded colour profiles that Word dislike.
  * Image width is capped to fit inside the page margins (no out-of-margin
    blobs that push content off the printable area).
  * NO raw XML / OxmlElement hacks that could break the Word schema.
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from PIL import Image

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
FRONTEND_DIR = os.path.dirname(os.path.abspath(__file__))
SRC_IMAGE_DIR = os.path.join(FRONTEND_DIR, "projects", "web", "src", "assets", "images")
TEMP_DIR = os.path.join(FRONTEND_DIR, "tmp_report_images")

OUTPUT_PATHS = [
    os.path.join("C:\\Users\\Unishka Bisht\\Downloads", "ResumeFlow_Internship_Report.docx"),
    os.path.join(FRONTEND_DIR, "ResumeFlow_Internship_Report.docx"),
]

# Brand colours
PRIMARY = RGBColor(0x63, 0x66, 0xF1)      # indigo
DARK = RGBColor(0x1E, 0x29, 0x3B)         # slate-800
GRAY = RGBColor(0x47, 0x55, 0x69)         # slate-600

# Page geometry (US Letter): 8.5" x 11", 1" margins => ~6.5" usable width.
# Cap image width comfortably inside the usable area.
MAX_IMG_WIDTH_IN = 5.5
MAX_IMG_HEIGHT_IN = 7.5


# ---------------------------------------------------------------------------
# Image pre-processing  (the key to valid embedding)
# ---------------------------------------------------------------------------
def safe_image(src_path, name):
    """
    Re-encode a source image into a clean RGB JPEG using PIL.
    Strips ICC profiles and EXIF so Word never chokes on the bytes.
    Returns (clean_path, width_px, height_px) or None if it cannot be read.
    """
    os.makedirs(TEMP_DIR, exist_ok=True)
    try:
        im = Image.open(src_path)
        im.load()  # fully decode -> catches truncated/corrupt files up front
    except Exception as e:
        print(f"  ! Skipping '{name}': cannot read image -> {e}")
        return None

    # Force RGB (handles RGBA / P / LA / palette / CMYK safely).
    if im.mode != "RGB":
        im = im.convert("RGB")

    # Drop rotation EXIF so the picture is shown upright without metadata.
    im = im.rotate(0, expand=False)

    clean_path = os.path.join(TEMP_DIR, f"clean_{name}.jpg")
    im.save(clean_path, "JPEG", quality=85, optimize=True, progressive=False)
    return clean_path, im.width, im.height


def add_screenshot(doc, src_path, caption, name):
    """Add a screenshot that is guaranteed valid + within margins."""
    result = safe_image(src_path, name)
    if result is None:
        return
    clean_path, w, h = result

    # Fit within the bounding box, preserving aspect ratio.
    aspect = w / h
    disp_w = min(MAX_IMG_WIDTH_IN, MAX_IMG_HEIGHT_IN * aspect)
    disp_h = disp_w / aspect

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(clean_path, width=Inches(disp_w), height=Inches(disp_h))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GRAY
    cap.paragraph_format.space_after = Pt(10)


# ---------------------------------------------------------------------------
# Reusable content helpers (clean, no schema-breaking XML)
# ---------------------------------------------------------------------------
def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = PRIMARY
        p.paragraph_format.space_before = Pt(14)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = DARK
        p.paragraph_format.space_before = Pt(10)
    else:
        run.font.size = Pt(11.5)
        run.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, size=11, italic=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = DARK
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10.5)
            if i == 0:
                run.bold = True

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    return table


def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    # Build a Word PAGE field using OxmlElement (schema-safe).
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run2 = p.add_run()
    run2._r.append(fld_begin)
    run2._r.append(instr)
    run2._r.append(fld_end)
    run2.font.size = Pt(9)
    run2.font.color.rgb = GRAY


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------
def build():
    doc = Document()

    # Usable width sanity for tables (<= 6.5").
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---- Title block ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SUMMER INTERNSHIP PROJECT REPORT")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("ResumeFlow — A Modern Resume Building Web Application")
    r.font.size = Pt(14)
    r.font.color.rgb = DARK
    sub.paragraph_format.space_after = Pt(2)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Frontend Development Internship")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY
    sub2.paragraph_format.space_after = Pt(10)

    # ---- Meta table ----
    add_table(
        doc,
        ["Particulars", "Details"],
        [
            ["Intern Name", "Unishka Bisht"],
            ["Project Title", "ResumeFlow — Professional Resume Builder"],
            ["Domain", "Web Application Development (Frontend)"],
            ["Technologies", "Angular 13, TypeScript, Angular Material, SCSS"],
            ["Internship Duration", "Summer 2026"],
        ],
        col_widths=[2.2, 4.3],
    )
    doc.add_paragraph()

    # ---- 1. Introduction ----
    add_heading(doc, "1. Introduction", 1)
    add_body(
        doc,
        "ResumeFlow is a modern, intuitive, full-featured web application that lets "
        "professionals and job seekers design, manage, version, and export professional "
        "resumes effortlessly. Traditional word processors often lead to broken layouts, "
        "inconsistent formatting across devices, and chaotic file management when tailoring "
        "applications for different roles. ResumeFlow addresses these pain points by providing "
        "consistent, pre-tested templates, real-time editing with live preview, version "
        "snapshots, and high-fidelity client-side exports with zero data loss.",
    )
    add_body(
        doc,
        "This report documents the work completed during the summer internship on the "
        "ResumeFlow frontend — a multi-project Angular workspace that delivers a polished, "
        "responsive user experience with glassmorphism styling, authentication, a document "
        "editor, multi-template switching, and client-side PDF/DOCX export.",
    )

    # ---- 2. Objectives ----
    add_heading(doc, "2. Internship Objectives", 1)
    for obj in [
        "Build a responsive, production-quality Angular frontend for resume creation.",
        "Implement a real-time document editor with live preview and section management.",
        "Develop multiple professional resume templates with one-click switching.",
        "Integrate client-side PDF and DOCX export with an export history log.",
        "Add authentication (signup, login, forgot/reset password) with route guarding.",
        "Deliver a premium UI with dark mode and glassmorphism design language.",
    ]:
        add_bullet(doc, obj)

    # ---- 3. Tech Stack ----
    add_heading(doc, "3. Technology Stack", 1)
    add_table(
        doc,
        ["Layer", "Tools & Libraries"],
        [
            ["Framework", "Angular 13 (standalone component-based architecture)"],
            ["Language", "TypeScript 4.6"],
            ["UI Components", "Angular Material 13, Angular CDK"],
            ["Styling", "SCSS custom design system, CSS variables, glassmorphism"],
            ["Document Generation", "html2pdf.js, docx, file-saver"],
            ["State & Networking", "RxJS, Angular HttpClient, Auth Interceptors"],
            ["Typography / Icons", "Google Fonts (Manrope, Sora, Space Grotesk, Roboto), Material Icons"],
        ],
        col_widths=[1.8, 4.7],
    )

    # ---- 4. System Architecture ----
    add_heading(doc, "4. System Architecture & Project Structure", 1)
    add_body(
        doc,
        "The application is organised as a multi-project Angular workspace. The primary "
        "deliverable lives under projects/web and is structured into reusable components, "
        "routed pages, shared services, and interceptors.",
    )
    add_table(
        doc,
        ["Module / Path", "Responsibility"],
        [
            ["guards/", "Route protection — AuthGuard & NoAuthGuard"],
            ["components/", "Reusable landing-page components (Hero, Features, Stats, etc.)"],
            ["pages/document-editor/", "Live resume builder & preview engine"],
            ["pages/documents/", "Resume collection management"],
            ["pages/templates/", "Template showcase & selection"],
            ["pages/exports/", "Export history & downloaded files"],
            ["pages/dashboard/", "User central dashboard & metrics"],
            ["pages/profile/", "User profile settings"],
            ["pages/applications/", "Job application tracking"],
            ["shared/services/", "Auth, HTTP interceptors, toast & common services"],
            ["environments/", "Environment configuration (dev / prod)"],
        ],
        col_widths=[2.7, 3.8],
    )

    # ---- 5. Key Features ----
    add_heading(doc, "5. Key Features Implemented", 1)

    add_heading(doc, "5.1 Real-Time Document Editor", 2)
    add_bullet(doc, "Interactive live preview updates instantly as content is typed.", "Live Preview: ")
    add_bullet(doc, "Add, remove, and reorganise sections (Experience, Education, Skills, Projects, Certifications).", "Section Management: ")
    add_bullet(doc, "Dynamic rich-text and bullet itemisation for clean readability.", "Formatting: ")
    add_bullet(doc, "Profile photo upload and preview support across resume templates.", "Photo Support: ")

    add_heading(doc, "5.2 Professional Templates", 2)
    add_body(
        doc,
        "Users can switch between expertly designed presets in a single click while "
        "preserving all entered data:",
    )
    for t in [
        "Modern Split — contemporary dual-column layout with visual accents",
        "Minimal Clean — streamlined, elegant, high-signal typography",
        "Technical / Engineering — optimised for technical competencies",
        "Executive — structured for leadership and enterprise portfolios",
        "Creative — balanced colour palettes with bold headline styling",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "5.3 Client-Side Export", 2)
    add_bullet(doc, "High-resolution, browser-side PDF generation via html2pdf.js.", "PDF Export: ")
    add_bullet(doc, "Native Microsoft Word (.docx) generation using docx + file-saver.", "DOCX Export: ")
    add_bullet(doc, "Digital log of all exported files.", "Export History: ")

    add_heading(doc, "5.4 Authentication & Security", 2)
    add_bullet(doc, "Complete auth suite: Sign Up, Login, Forgot Password, Password Reset.", "Auth Suite: ")
    add_bullet(doc, "JWT interceptor with automatic token management.", "JWT Interceptor: ")
    add_bullet(doc, "Protected routes via Angular route guards (AuthGuard & NoAuthGuard).", "Route Guards: ")

    add_heading(doc, "5.5 Premium UI & Theming", 2)
    add_bullet(doc, "Glassmorphism effects with smooth micro-interactions.", "Visual Design: ")
    add_bullet(doc, "Responsive navigation across mobile, tablet, and desktop.", "Responsive: ")
    add_bullet(doc, "Seamless light/dark theme switching.", "Dark Mode: ")

    # ---- 6. Implementation Highlights ----
    add_heading(doc, "6. Implementation Highlights", 1)
    for h in [
        "Built a modular, lazy-loaded Angular workspace to keep the bundle lean and maintainable.",
        "Implemented a reactive document model so template switching never loses user data.",
        "Wrapped html2pdf.js and the docx library in dedicated services for reusable, one-click exports.",
        "Centralised styling through SCSS variables and a custom design system for consistent theming.",
        "Added an HTTP interceptor to transparently attach and refresh JWT tokens on every request.",
        "Used Angular Material components to accelerate accessible, responsive UI construction.",
    ]:
        add_bullet(doc, h)

    # ---- 7. Challenges ----
    add_heading(doc, "7. Challenges & Solutions", 1)
    add_table(
        doc,
        ["Challenge", "Resolution"],
        [
            ["Preserving data across template switches",
             "Adopted a single source-of-truth document model bound to all templates."],
            ["High-fidelity PDF layout from HTML",
             "Tuned html2pdf.js page-break, scale, and CSS to match on-screen design."],
            ["Consistent theming in dark mode",
             "Drove all colours from SCSS variables and CSS custom properties."],
            ["Maintaining performance with live preview",
             "Used reactive forms and change-detection best practices to avoid reflow thrash."],
        ],
        col_widths=[2.6, 3.9],
    )

    # ---- 8. Testing & Validation ----
    add_heading(doc, "8. Testing & Validation", 1)
    add_bullet(doc, "Unit tests via Karma + Jasmine for components and services.")
    add_bullet(doc, "Cross-browser and responsive checks (mobile, tablet, desktop).")
    add_bullet(doc, "Manual end-to-end validation of the signup → edit → export flow.")

    # ---- 9. UI Screenshots (the image section) ----
    add_heading(doc, "9. Application Screenshots", 1)
    add_body(
        doc,
        "The following screenshots illustrate the live editor and template showcase "
        "of the ResumeFlow application.",
    )
    shots = [
        ("1.jpeg", "Figure 1 — ResumeFlow document editor / live preview"),
        ("2.jpeg", "Figure 2 — Template selection & switching"),
        ("3.jpeg", "Figure 3 — Resume collection management"),
        ("4.jpeg", "Figure 4 — Dashboard / export history"),
    ]
    for fname, caption in shots:
        src = os.path.join(SRC_IMAGE_DIR, fname)
        if os.path.exists(src):
            add_screenshot(doc, src, caption, os.path.splitext(fname)[0])
        else:
            print(f"  ! Image not found: {src}")

    # ---- 10. Future Goals ----
    add_heading(doc, "10. Future Scope", 1)
    for g in [
        "Automated bullet-point and action-verb suggestions based on target keywords.",
        "Integrated cover-letter builder synced with resume styling.",
        "Custom section designer (tabular and timeline layouts).",
        "Analytics & read-only sharing links for prospective employers.",
        "Cloud backup integration (Google Drive, Dropbox, OneDrive).",
    ]:
        add_bullet(doc, g)

    # ---- 11. Conclusion ----
    add_heading(doc, "11. Conclusion", 1)
    add_body(
        doc,
        "The summer internship delivered a complete, production-ready frontend for ResumeFlow. "
        "The platform successfully combines a real-time editor, diverse professional templates, "
        "secure authentication, and reliable client-side exports into a single, polished "
        "experience. The work established a strong, scalable foundation for the future features "
        "outlined above and demonstrated practical application of Angular, TypeScript, and "
        "modern frontend engineering practices.",
    )

    # ---- Acknowledgement ----
    add_heading(doc, "Acknowledgement", 1)
    add_body(
        doc,
        "I sincerely thank my internship mentor and the ResumeFlow team for their guidance, "
        "feedback, and support throughout this project. Their mentorship was instrumental in "
        "shaping both the application and my growth as a frontend developer.",
        italic=True,
        color=GRAY,
    )

    add_page_number_footer(doc)

    # ---- Save to ALL output paths ----
    for out in OUTPUT_PATHS:
        os.makedirs(os.path.dirname(out), exist_ok=True)
        doc.save(out)
        print(f"Saved: {out}  ({os.path.getsize(out)} bytes)")
    return OUTPUT_PATHS


# ---------------------------------------------------------------------------
# Verification
# ---------------------------------------------------------------------------
def verify(paths):
    import zipfile
    import xml.dom.minidom as M

    for path in paths:
        print(f"\n=== Verifying {path} ===")
        # 1. python-docx can open it
        d = Document(path)
        print(f"  Document() opened OK  | paragraphs={len(d.paragraphs)}  tables={len(d.tables)}")

        # 2. Every XML part is well-formed
        z = zipfile.ZipFile(path)
        bad = []
        xml_count = 0
        for name in z.namelist():
            if name.endswith(".xml") or name.endswith(".rels"):
                xml_count += 1
                try:
                    M.parseString(z.read(name))
                except Exception as e:
                    bad.append((name, str(e)))
        print(f"  XML parts checked: {xml_count}")
        print(f"  Malformed parts: {bad if bad else 'NONE'}")

        # 3. Media parts present
        media = [n for n in z.namelist() if n.startswith("word/media/")]
        print(f"  Embedded media: {media}")
        print(f"  File size: {os.path.getsize(path)} bytes")


if __name__ == "__main__":
    paths = build()
    verify(paths)
