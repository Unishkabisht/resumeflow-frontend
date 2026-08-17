"""
build_report.py
Generates the Summer Internship Project Report for the ResumeFlow
web application and writes it to:

    C:\\Users\\Unishka Bisht\\resumeflow-frontend\\Summer_Internship_Project_Report_ResumeFlow.docx

Run with:  python build_report.py
"""

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

# ---------------------------------------------------------------------------
# Output path
# ---------------------------------------------------------------------------
OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Summer_Internship_Project_Report_ResumeFlow.docx",
)

# Brand colours
PRIMARY = RGBColor(0x63, 0x66, 0xF1)      # indigo
DARK = RGBColor(0x1E, 0x29, 0x3B)         # slate-800
GRAY = RGBColor(0x47, 0x55, 0x69)         # slate-600
ACCENT = RGBColor(0xFF, 0x40, 0x81)       # material pink


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------
def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    """Add a styled heading paragraph."""
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = PRIMARY
        # bottom border
        p_pr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "C7D2FE")
        pbdr.append(bottom)
        p_pr.append(pbdr)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = DARK
    else:
        run.font.size = Pt(11.5)
        run.font.color.rgb = GRAY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, size=11, italic=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = DARK
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], "6366F1")

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10.5)
            if i == 0:
                run.bold = True

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    return table


def _add_field(paragraph, instr):
    """Insert a Word field (e.g. PAGE) as a run-level field."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)
    return run


def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    r2 = _add_field(p, "PAGE")
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRAY


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------
def build():
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---- Title block ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SUMMER INTERNSHIP PROJECT REPORT")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("ResumeFlow — A Modern Resume Building Web Application")
    r.font.size = Pt(14)
    r.font.color.rgb = DARK
    sub.paragraph_format.space_after = Pt(2)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Frontend Development Internship")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY
    sub2.paragraph_format.space_after = Pt(10)

    # ---- Meta table ----
    meta = add_table(
        doc,
        ["Particulars", "Details"],
        [
            ["Intern Name", "Unishka Bisht"],
            ["Project Title", "ResumeFlow — Professional Resume Builder"],
            ["Domain", "Web Application Development (Frontend)"],
            ["Technologies", "Angular 13, TypeScript, Angular Material, SCSS"],
            ["Internship Duration", "Summer 2026"],
            ["Report Generated", os.path.basename(OUTPUT_PATH)],
        ],
        col_widths=[2.2, 4.3],
    )
    doc.add_paragraph()

    # ---- 1. Introduction ----
    add_heading(doc, "1. Introduction", 1)
    add_body(
        doc,
        "ResumeFlow is a modern, intuitive, full-featured web application that lets "
        "professionals and job seekers design, manage, version, and export professional "
        "resumes effortlessly. Traditional word processors often lead to broken layouts, "
        "inconsistent formatting across devices, and chaotic file management when tailoring "
        "applications for different roles. ResumeFlow addresses these pain points by providing "
        "consistent, pre-tested templates, real-time editing with live preview, version "
        "snapshots, and high-fidelity client-side exports with zero data loss.",
    )
    add_body(
        doc,
        "This report documents the work completed during the summer internship on the "
        "ResumeFlow frontend — a multi-project Angular workspace that delivers a polished, "
        "responsive user experience with glassmorphism styling, authentication, a document "
        "editor, multi-template switching, and client-side PDF/DOCX export.",
    )

    # ---- 2. Objectives ----
    add_heading(doc, "2. Internship Objectives", 1)
    for obj in [
        "Build a responsive, production-quality Angular frontend for resume creation.",
        "Implement a real-time document editor with live preview and section management.",
        "Develop multiple professional resume templates with one-click switching.",
        "Integrate client-side PDF and DOCX export with an export history log.",
        "Add authentication (signup, login, forgot/reset password) with route guarding.",
        "Deliver a premium UI with dark mode and glassmorphism design language.",
    ]:
        add_bullet(doc, obj)

    # ---- 3. Tech Stack ----
    add_heading(doc, "3. Technology Stack", 1)
    add_table(
        doc,
        ["Layer", "Tools & Libraries"],
        [
            ["Framework", "Angular 13 (standalone component-based architecture)"],
            ["Language", "TypeScript 4.6"],
            ["UI Components", "Angular Material 13, Angular CDK"],
            ["Styling", "SCSS custom design system, CSS variables, glassmorphism"],
            ["Document Generation", "html2pdf.js, docx, file-saver"],
            ["State & Networking", "RxJS, Angular HttpClient, Auth Interceptors"],
            ["Typography / Icons", "Google Fonts (Manrope, Sora, Space Grotesk, Roboto), Material Icons"],
        ],
        col_widths=[1.8, 4.7],
    )

    # ---- 4. System Architecture ----
    add_heading(doc, "4. System Architecture & Project Structure", 1)
    add_body(
        doc,
        "The application is organised as a multi-project Angular workspace. The primary "
        "deliverable lives under projects/web and is structured into reusable components, "
        "routed pages, shared services, and interceptors.",
    )
    add_table(
        doc,
        ["Module / Path", "Responsibility"],
        [
            ["guards/", "Route protection — AuthGuard & NoAuthGuard"],
            ["components/", "Reusable landing-page components (Hero, Features, Stats, etc.)"],
            ["pages/document-editor/", "Live resume builder & preview engine"],
            ["pages/documents/", "Resume collection management"],
            ["pages/templates/", "Template showcase & selection"],
            ["pages/exports/", "Export history & downloaded files"],
            ["pages/dashboard/", "User central dashboard & metrics"],
            ["pages/profile/", "User profile settings"],
            ["pages/applications/", "Job application tracking"],
            ["shared/services/", "Auth, HTTP interceptors, toast & common services"],
            ["environments/", "Environment configuration (dev / prod)"],
        ],
        col_widths=[2.7, 3.8],
    )

    # ---- 5. Key Features ----
    add_heading(doc, "5. Key Features Implemented", 1)

    add_heading(doc, "5.1 Real-Time Document Editor", 2)
    add_bullet(doc, "Interactive live preview updates instantly as content is typed.", "Live Preview: ")
    add_bullet(doc, "Add, remove, and reorganise sections (Experience, Education, Skills, Projects, Certifications).", "Section Management: ")
    add_bullet(doc, "Dynamic rich-text and bullet itemisation for clean readability.", "Formatting: ")
    add_bullet(doc, "Profile photo upload and preview support across resume templates.", "Photo Support: ")

    add_heading(doc, "5.2 Professional Templates", 2)
    add_body(
        doc,
        "Users can switch between expertly designed presets in a single click while "
        "preserving all entered data:",
    )
    for t in [
        "Modern Split — contemporary dual-column layout with visual accents",
        "Minimal Clean — streamlined, elegant, high-signal typography",
        "Technical / Engineering — optimised for technical competencies",
        "Executive — structured for leadership and enterprise portfolios",
        "Creative — balanced colour palettes with bold headline styling",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "5.3 Client-Side Export", 2)
    add_bullet(doc, "High-resolution, browser-side PDF generation via html2pdf.js.", "PDF Export: ")
    add_bullet(doc, "Native Microsoft Word (.docx) generation using docx + file-saver.", "DOCX Export: ")
    add_bullet(doc, "Digital log of all exported files.", "Export History: ")

    add_heading(doc, "5.4 Authentication & Security", 2)
    add_bullet(doc, "Complete auth suite: Sign Up, Login, Forgot Password, Password Reset.", "Auth Suite: ")
    add_bullet(doc, "JWT interceptor with automatic token management.", "JWT Interceptor: ")
    add_bullet(doc, "Protected routes via Angular route guards (AuthGuard & NoAuthGuard).", "Route Guards: ")

    add_heading(doc, "5.5 Premium UI & Theming", 2)
    add_bullet(doc, "Glassmorphism effects with smooth micro-interactions.", "Visual Design: ")
    add_bullet(doc, "Responsive navigation across mobile, tablet, and desktop.", "Responsive: ")
    add_bullet(doc, "Seamless light/dark theme switching.", "Dark Mode: ")

    # ---- 6. Implementation Highlights ----
    add_heading(doc, "6. Implementation Highlights", 1)
    for h in [
        "Built a modular, lazy-loaded Angular workspace to keep the bundle lean and maintainable.",
        "Implemented a reactive document model so template switching never loses user data.",
        "Wrapped html2pdf.js and the docx library in dedicated services for reusable, one-click exports.",
        "Centralised styling through SCSS variables and a custom design system for consistent theming.",
        "Added an HTTP interceptor to transparently attach and refresh JWT tokens on every request.",
        "Used Angular Material components to accelerate accessible, responsive UI construction.",
    ]:
        add_bullet(doc, h)

    # ---- 7. Challenges ----
    add_heading(doc, "7. Challenges & Solutions", 1)
    add_table(
        doc,
        ["Challenge", "Resolution"],
        [
            ["Preserving data across template switches",
             "Adopted a single source-of-truth document model bound to all templates."],
            ["High-fidelity PDF layout from HTML",
             "Tuned html2pdf.js page-break, scale, and CSS to match on-screen design."],
            ["Consistent theming in dark mode",
             "Drove all colours from SCSS variables and CSS custom properties."],
            ["Maintaining performance with live preview",
             "Used reactive forms and change-detection best practices to avoid reflow thrash."],
        ],
        col_widths=[2.6, 3.9],
    )

    # ---- 8. Testing & Validation ----
    add_heading(doc, "8. Testing & Validation", 1)
    add_bullet(doc, "Unit tests via Karma + Jasmine for components and services.")
    add_bullet(doc, "Cross-browser and responsive checks (mobile, tablet, desktop).")
    add_bullet(doc, "Manual end-to-end validation of the signup → edit → export flow.")

    # ---- 9. Future Goals ----
    add_heading(doc, "9. Future Scope", 1)
    for g in [
        "Automated bullet-point and action-verb suggestions based on target keywords.",
        "Integrated cover-letter builder synced with resume styling.",
        "Custom section designer (tabular and timeline layouts).",
        "Analytics & read-only sharing links for prospective employers.",
        "Cloud backup integration (Google Drive, Dropbox, OneDrive).",
    ]:
        add_bullet(doc, g)

    # ---- 10. Conclusion ----
    add_heading(doc, "10. Conclusion", 1)
    add_body(
        doc,
        "The summer internship delivered a complete, production-ready frontend for ResumeFlow. "
        "The platform successfully combines a real-time editor, diverse professional templates, "
        "secure authentication, and reliable client-side exports into a single, polished "
        "experience. The work established a strong, scalable foundation for the future features "
        "outlined above and demonstrated practical application of Angular, TypeScript, and "
        "modern frontend engineering practices.",
    )

    # ---- Acknowledgement ----
    add_heading(doc, "Acknowledgement", 1)
    add_body(
        doc,
        "I sincerely thank my internship mentor and the ResumeFlow team for their guidance, "
        "feedback, and support throughout this project. Their mentorship was instrumental in "
        "shaping both the application and my growth as a frontend developer.",
        italic=True,
        color=GRAY,
    )

    add_page_number_footer(doc)

    # ---- Save ----
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build()
    print("File exists:", os.path.exists(path))
    print("File size (bytes):", os.path.getsize(path) if os.path.exists(path) else 0)
